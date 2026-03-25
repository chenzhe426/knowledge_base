"""
DOCX parser with full structure preservation.

Each run of the document produces a list of typed blocks (heading,
paragraph, list_item, table, quote, code) with style metadata.
The unified ParsedDocument is built from the best-performing strategy.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from app.ingestion.config import CleaningConfig, DocxParserConfig, ParsingConfig
from app.ingestion.normalizers import (
    blocks_to_content,
    clean_blocks,
    normalize_text,
)
from app.ingestion.parsers.base import BaseParser
from app.ingestion.schemas import ParsedDocument

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Section path helper
# ---------------------------------------------------------------------------


def _section_path_push(section_stack: list[str], heading_text: str, level: int) -> str:
    while len(section_stack) >= level:
        section_stack.pop()
    section_stack.append(heading_text)
    return " > ".join(section_stack)


# ---------------------------------------------------------------------------
# Strategy 1 – docling (best for complex Word docs)
# ---------------------------------------------------------------------------


def _parse_docx_with_docling(path: Path) -> list[dict[str, Any]]:
    """Use docling's DOCX support for structured extraction."""
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(path))
    markdown_text = result.document.export_to_markdown()

    lines = normalize_text(markdown_text).split("\n")
    blocks: list[dict[str, Any]] = []
    section_stack: list[str] = []
    para_buf: list[str] = []

    def flush_paragraph():
        nonlocal para_buf
        if para_buf:
            text = "\n".join(para_buf).strip()
            if text:
                blocks.append(
                    {
                        "type": "paragraph",
                        "text": text,
                        "section_path": " > ".join(section_stack) if section_stack else None,
                        "parser_source": "docling",
                    }
                )
            para_buf = []

    for line in lines:
        s = line.strip()
        if not s:
            flush_paragraph()
            continue
        if s.startswith("#"):
            flush_paragraph()
            level = len(s) - len(s.lstrip("#"))
            heading_text = s[level:].strip()
            section_path = _section_path_push(section_stack, heading_text, level)
            blocks.append(
                {
                    "type": "heading",
                    "text": heading_text,
                    "heading_level": level,
                    "section_path": section_path,
                    "parser_source": "docling",
                }
            )
            continue
        if s.startswith("|") and s.endswith("|"):
            flush_paragraph()
            blocks.append(
                {
                    "type": "table",
                    "text": s,
                    "section_path": " > ".join(section_stack) if section_stack else None,
                    "parser_source": "docling",
                }
            )
            continue
        if re.match(r"^\s*([-*•·]|\d+[.)、])\s+.+$", s):
            flush_paragraph()
            blocks.append(
                {
                    "type": "list_item",
                    "text": s,
                    "section_path": " > ".join(section_stack) if section_stack else None,
                    "parser_source": "docling",
                }
            )
            continue
        para_buf.append(s)

    flush_paragraph()
    return blocks


# ---------------------------------------------------------------------------
# Strategy 2 – python-docx (reliable fallback)
# ---------------------------------------------------------------------------


def _parse_docx_with_python_docx(
    path: Path,
    config: DocxParserConfig,
) -> list[dict[str, Any]]:
    """
    Extract blocks from DOCX using python-docx.

    Style mapping
    -------------
    python-docx exposes paragraph styles via ``para.style.name``.
    We map those to our canonical block types and heading levels.
    """
    from docx import Document

    doc = Document(str(path))
    blocks: list[dict[str, Any]] = []
    section_stack: list[str] = []

    # ---- Document properties (optional) ----
    core_props: dict[str, str] = {}
    if config.extract_doc_props:
        try:
            cp = doc.core_properties
            core_props = {
                k: str(v) for k, v in {
                    "title": cp.title,
                    "author": cp.author,
                    "subject": cp.subject,
                    "keywords": cp.keywords,
                }.items() if v
            }
        except Exception:
            pass

    # ---- Extract paragraphs ----
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue

        style_name = ""
        try:
            style_name = (para.style.name or "").lower().strip()
        except Exception:
            style_name = ""

        bold = False
        italic = False
        try:
            for run in para.runs:
                if run.bold:
                    bold = True
                if run.italic:
                    italic = True
        except Exception:
            pass

        style_meta = {
            "style_name": para.style.name if para.style else "",
            "bold": bold,
            "italic": italic,
        }

        # ---- Heading detection via style ----
        if style_name in config.style_to_level:
            level = config.style_to_level[style_name]
            section_path = _section_path_push(section_stack, text, level)
            blocks.append(
                {
                    "type": "heading",
                    "text": text,
                    "heading_level": level,
                    "section_path": section_path,
                    "block_index": i,
                    "parser_source": "python-docx",
                    "style_meta": style_meta,
                }
            )
            continue

        # Fallback: all-caps short text as potential heading
        allcaps_ratio = sum(1 for c in text if c.isupper()) / max(len(text), 1)
        if allcaps_ratio > 0.8 and len(text) < 100 and len(text) > 2:
            level = 1
            section_path = _section_path_push(section_stack, text, level)
            blocks.append(
                {
                    "type": "heading",
                    "text": text,
                    "heading_level": level,
                    "section_path": section_path,
                    "block_index": i,
                    "parser_source": "python-docx",
                    "style_meta": style_meta,
                }
            )
            continue

        # ---- List item detection ----
        if re.match(config.list_pattern, text):
            blocks.append(
                {
                    "type": "list_item",
                    "text": text,
                    "section_path": " > ".join(section_stack) if section_stack else None,
                    "block_index": i,
                    "parser_source": "python-docx",
                    "style_meta": style_meta,
                }
            )
            continue

        # ---- Quote / blockquote ----
        if style_name in ("quote", "blockquote", "citation"):
            blocks.append(
                {
                    "type": "quote",
                    "text": text,
                    "section_path": " > ".join(section_stack) if section_stack else None,
                    "block_index": i,
                    "parser_source": "python-docx",
                    "style_meta": style_meta,
                }
            )
            continue

        # ---- Regular paragraph ----
        blocks.append(
            {
                "type": "paragraph",
                "text": text,
                "section_path": " > ".join(section_stack) if section_stack else None,
                "block_index": i,
                "parser_source": "python-docx",
                "style_meta": style_meta,
            }
        )

    # ---- Extract tables ----
    if config.extract_tables:
        for i, table in enumerate(doc.tables):
            table_lines: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
            if table_lines:
                blocks.append(
                    {
                        "type": "table",
                        "text": "\n".join(table_lines),
                        "section_path": " > ".join(section_stack) if section_stack else None,
                        "block_index": len(doc.paragraphs) + i,
                        "parser_source": "python-docx",
                    }
                )

    return blocks


# ---------------------------------------------------------------------------
# Quality scoring for DOCX (simpler than PDF)
# ---------------------------------------------------------------------------


def _score_docx_blocks(blocks: list[dict[str, Any]]) -> float:
    """Simple heuristic quality score for DOCX blocks."""
    if not blocks:
        return 0.0

    texts = [b.get("text", "") or "" for b in blocks]
    total_chars = sum(len(t) for t in texts)
    nonempty = sum(1 for t in texts if t.strip())
    heading_count = sum(1 for b in blocks if b.get("type") == "heading")
    table_count = sum(1 for b in blocks if b.get("type") == "table")
    list_count = sum(1 for b in blocks if b.get("type") == "list_item")

    score = (
        0.20 * min(total_chars / 3000, 1.0)
        + 0.20 * min(nonempty / 30, 1.0)
        + 0.20 * min(heading_count / 5, 1.0)
        + 0.20 * (1.0 if table_count > 0 else 0.0)
        + 0.20 * (1.0 if list_count > 0 else 0.0)
    )
    return max(0.0, min(1.0, score))


# ---------------------------------------------------------------------------
# DOCX Parser – public interface
# ---------------------------------------------------------------------------


class DocxParser(BaseParser):
    """
    Multi-strategy DOCX parser.

    Runs docling (preferred) and python-docx (fallback), scores each result,
    picks the best, and returns a fully-structured ParsedDocument.
    """

    file_type = "docx"

    def __init__(self, config: ParsingConfig | None = None):
        self._cfg = config.docx if config else DocxParserConfig()
        self._clean_cfg = config.cleaning if config else CleaningConfig()

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path)
        candidates: dict[str, Any] = {}
        all_errors: list[str] = []

        strategies: list[tuple[str, Any]] = [
            ("docling", _parse_docx_with_docling),
        ]
        strategies.append(
            ("python-docx", lambda p: _parse_docx_with_python_docx(p, self._cfg))
        )

        for name, fn in strategies:
            try:
                blocks = fn(path)
                if not blocks:
                    all_errors.append(f"{name}: no blocks returned")
                    continue
                score = _score_docx_blocks(blocks)
                candidates[name] = {"blocks": blocks, "quality_score": score}
            except Exception as e:
                all_errors.append(f"{name}: {e}")

        if not candidates:
            raise RuntimeError(
                f"all DOCX parsers failed for {path}. Errors: {all_errors}"
            )

        # Select best
        best_name = max(candidates, key=lambda k: candidates[k]["quality_score"])
        best = candidates[best_name]
        blocks = best["blocks"]

        if self._cfg.verbose:
            logger.info(
                "[DocxParser] strategy scores – %s",
                {k: v["quality_score"] for k, v in candidates.items()},
            )

        # Tag blocks
        for block in blocks:
            block["source_format"] = "docx"

        cleaned_blocks = clean_blocks(blocks, self._clean_cfg)
        for b in cleaned_blocks:
            b.setdefault("source_format", "docx")

        content = blocks_to_content(
            cleaned_blocks,
            include_headings=self._clean_cfg.include_headings_in_content,
        )

        table_count = sum(1 for b in cleaned_blocks if b.get("type") == "table")
        heading_count = sum(1 for b in cleaned_blocks if b.get("type") == "heading")

        metadata = {
            "parser_used": best_name,
            "candidate_scores": {k: v["quality_score"] for k, v in candidates.items()},
            "selection_reason": f"score={best['quality_score']:.3f}",
            "quality_score": best["quality_score"],
            "table_count": table_count,
            "heading_count": heading_count,
            "source_format": "docx",
            "source_file": str(path),
            "errors": all_errors,
        }

        return ParsedDocument(
            title=path.stem,
            content=content,
            blocks=cleaned_blocks,
            source_path=str(path),
            file_type=self.file_type,
            metadata=metadata,
        )
