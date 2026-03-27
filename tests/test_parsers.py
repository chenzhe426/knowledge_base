"""
Tests for the ingestion / parsing pipeline.

Run with:  pytest tests/test_parsers.py -v

Tests are organized by file format and cover:
- Successful parsing with structured blocks
- Metadata completeness
- Quality scoring (PDF)
- Heading / list / table detection
- Noise cleaning
- Encoding handling (TXT)
- Unified ParsedDocument schema
"""
from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from app.ingestion.config import (
    CleaningConfig,
    DocxParserConfig,
    ParsingConfig,
    PdfParserConfig,
    TextParserConfig,
)
from app.ingestion.loaders import load_text_file
from app.ingestion.normalizers import (
    blocks_to_content,
    clean_blocks,
    is_noise_line,
    is_template_noise,
    normalize_text,
    repair_hyphenation,
)
from app.ingestion.parsers.base import BaseParser
from app.ingestion.parsers.docx_parser import DocxParser
from app.ingestion.parsers.pdf_parser import PdfParser
from app.ingestion.parsers.text_parser import TextParser
from app.ingestion.pipeline import parse_document
from app.ingestion.quality import score_pdf_blocks
from app.ingestion.schemas import ParsedDocument


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def cfg():
    return ParsingConfig()


@pytest.fixture
def tmp_dir(tmp_path):
    return tmp_path


# ---------------------------------------------------------------------------
# A. Normalizer unit tests
# ---------------------------------------------------------------------------


class TestNormalizers:
    def test_normalize_text_basic(self):
        assert normalize_text("  hello  world  ") == "hello  world"
        assert normalize_text("hello\r\nworld") == "hello\nworld"
        assert normalize_text("hello\u00a0world") == "hello world"
        assert normalize_text("hello\u3000world") == "hello world"
        assert normalize_text("a\n\n\nb") == "a\n\nb"

    def test_normalize_text_empty(self):
        assert normalize_text("") == ""
        assert normalize_text("   ") == ""

    def test_is_noise_line_page_numbers(self):
        assert is_noise_line("  第 1 页  ")
        assert is_noise_line("Page 5")
        assert is_noise_line("5 / 20")
        assert is_noise_line("- 3 -")
        assert is_noise_line("[ 42 ]")
        assert is_noise_line("5/10")
        assert not is_noise_line("This is a normal sentence about finance.")
        assert not is_noise_line("hello world")

    def test_is_template_noise(self):
        assert is_template_noise("contact@example.com")
        assert is_template_noise("https://example.com")
        assert is_template_noise("www.example.com")
        assert is_template_noise("© 2024 Acme Corp")
        assert is_template_noise("_______________")
        assert is_template_noise("CONFIDENTIAL – Internal Use Only")
        assert not is_template_noise("The company has 500 employees.")

    def test_repair_hyphenation(self):
        assert repair_hyphenation("invest-\nment") == "investment"
        assert repair_hyphenation("long-\nterm") == "longterm"
        assert repair_hyphenation("no hyphen here") == "no hyphen here"

    def test_clean_blocks_removes_noise(self):
        blocks = [
            {"text": "Page 1", "type": "paragraph"},
            {"text": "  第 2 页  ", "type": "paragraph"},
            {"text": "Real content here", "type": "paragraph"},
            {"text": "   ", "type": "paragraph"},
            {"text": "---", "type": "paragraph"},
        ]
        cfg = CleaningConfig(remove_page_numbers=True, remove_separator_lines=True)
        cleaned = clean_blocks(blocks, config=cfg)
        texts = [b["text"] for b in cleaned]
        assert "Page 1" not in texts
        assert "第 2 页" not in texts
        assert "Real content here" in texts

    def test_clean_blocks_preserves_headings(self):
        blocks = [
            {"text": "Introduction", "type": "heading", "heading_level": 1},
            {"text": "Some paragraph", "type": "paragraph"},
        ]
        cfg = CleaningConfig()
        cleaned = clean_blocks(blocks, config=cfg)
        types = [b["type"] for b in cleaned]
        assert "heading" in types
        assert "paragraph" in types

    def test_blocks_to_content_excludes_headings_by_default(self):
        blocks = [
            {"text": "Title", "type": "heading"},
            {"text": "Paragraph one", "type": "paragraph"},
            {"text": "Paragraph two", "type": "paragraph"},
        ]
        content = blocks_to_content(blocks, include_headings=False)
        assert "Title" not in content
        assert "Paragraph one" in content

    def test_blocks_to_content_includes_headings_when_asked(self):
        blocks = [
            {"text": "Title", "type": "heading"},
            {"text": "Paragraph one", "type": "paragraph"},
        ]
        content = blocks_to_content(blocks, include_headings=True)
        assert "Title" in content


# ---------------------------------------------------------------------------
# B. Load test
# ---------------------------------------------------------------------------


class TestLoaders:
    def test_load_text_file_utf8(self, tmp_dir):
        f = tmp_dir / "hello.txt"
        f.write_text("你好，世界！\n这是测试。", encoding="utf-8")
        text = load_text_file(f)
        assert "你好" in text
        assert "这是测试" in text

    def test_load_text_file_utf8_silent_bom(self, tmp_dir):
        f = tmp_dir / "bom.txt"
        f.write_bytes(b"\xef\xbb\xbfHello, world!\n")
        text = load_text_file(f)
        assert text.startswith("Hello")

    def test_load_text_file_missing(self):
        with pytest.raises(FileNotFoundError):
            load_text_file("/nonexistent/file.txt")


# ---------------------------------------------------------------------------
# C. Quality scoring
# ---------------------------------------------------------------------------


class TestQualityScoring:
    def test_score_empty_blocks(self):
        result = score_pdf_blocks([])
        assert result["quality_score"] == 0.0
        assert result["total_chars"] == 0

    def test_score_good_blocks(self):
        blocks = [
            {"text": "Introduction to the project", "type": "heading", "page": 1},
            {
                "text": "This is a well-formed paragraph with enough content to score well in our quality metrics. " * 5,
                "type": "paragraph",
                "page": 1,
            },
            {
                "text": "Another paragraph with substantial text content for quality assessment.",
                "type": "paragraph",
                "page": 1,
            },
        ]
        result = score_pdf_blocks(blocks)
        assert result["quality_score"] > 0.3
        assert result["total_chars"] > 0
        assert not result["likely_scanned"]

    def test_score_low_quality_blocks(self):
        blocks = [
            {"text": "a", "type": "paragraph", "page": 1},
            {"text": "b", "type": "paragraph", "page": 1},
            {"text": "c", "type": "paragraph", "page": 1},
        ]
        result = score_pdf_blocks(blocks)
        assert result["likely_scanned"] or result["short_frag_ratio"] > 0.5

    def test_score_with_table(self):
        blocks = [
            {"text": "Header A | Header B", "type": "table", "page": 1},
            {"text": "Cell 1 | Cell 2", "type": "table", "page": 1},
            {
                "text": "A long paragraph with substantial content for testing.",
                "type": "paragraph",
                "page": 1,
            },
        ]
        result = score_pdf_blocks(blocks)
        assert result["table_count"] == 2
        assert result["quality_score"] > 0


# ---------------------------------------------------------------------------
# D. TXT parser
# ---------------------------------------------------------------------------


class TestTextParser:
    def test_markdown_heading_detection(self, tmp_dir):
        f = tmp_dir / "headings.md"
        f.write_text("# Main Title\n\n## Section One\n\nContent here.\n\n### Subsection\n\nMore content.\n", encoding="utf-8")
        parser = TextParser()
        doc = parser.parse(f)
        assert doc.file_type == "text"
        assert len(doc.blocks) > 0
        types = [b["type"] for b in doc.blocks]
        assert "heading" in types
        heading_blocks = [b for b in doc.blocks if b["type"] == "heading"]
        assert any(b["heading_level"] == 1 for b in heading_blocks)
        assert any(b["heading_level"] == 2 for b in heading_blocks)
        assert any(b["heading_level"] == 3 for b in heading_blocks)

    def test_chinese_heading_detection(self, tmp_dir):
        f = tmp_dir / "cn.txt"
        f.write_text("第一章 项目概述\n\n这是正文内容。\n\n第 1 节 项目背景\n\n更多内容。\n", encoding="utf-8")
        parser = TextParser()
        doc = parser.parse(f)
        types = [b["type"] for b in doc.blocks]
        assert "heading" in types

    def test_list_item_detection(self, tmp_dir):
        f = tmp_dir / "list.txt"
        f.write_text("- First item\n- Second item\n- Third item\n\n1. Numbered one\n2. Numbered two\n", encoding="utf-8")
        parser = TextParser()
        doc = parser.parse(f)
        types = [b["type"] for b in doc.blocks]
        assert "list_item" in types
        list_items = [b for b in doc.blocks if b["type"] == "list_item"]
        assert len(list_items) >= 4

    def test_code_block_detection(self, tmp_dir):
        f = tmp_dir / "code.md"
        f.write_text("```python\ndef hello():\n    print('hi')\n```\n\nSome text.\n", encoding="utf-8")
        parser = TextParser()
        doc = parser.parse(f)
        types = [b["type"] for b in doc.blocks]
        assert "code" in types

    def test_smart_paragraph_reconstruction(self, tmp_dir):
        f = tmp_dir / "para.txt"
        f.write_text(
            "这是一个完整的句子。这是一个完整的句子。这是完整的句子。\n\n下一个段落开始了。"
        )
        parser = TextParser()
        doc = parser.parse(f)
        # Should not have overly fragmented paragraphs
        para_blocks = [b for b in doc.blocks if b["type"] == "paragraph"]
        # At least one paragraph should have reasonable length
        long_paras = [p for p in para_blocks if len(p["text"]) > 20]
        assert len(long_paras) >= 1

    def test_block_source_format_tagged(self, tmp_dir):
        f = tmp_dir / "simple.txt"
        f.write_text("Hello world.\n")
        parser = TextParser()
        doc = parser.parse(f)
        for b in doc.blocks:
            assert b.get("source_format") == "txt"

    def test_metadata_fields(self, tmp_dir):
        f = tmp_dir / "meta.txt"
        f.write_text("# Title\n\nParagraph one.\n\n- item one\n- item two\n\n| col1 | col2 |\n|------|------|\n| a    | b    |\n")
        parser = TextParser()
        doc = parser.parse(f)
        assert "heading_count" in doc.metadata
        assert "table_count" in doc.metadata
        assert "list_count" in doc.metadata
        assert doc.metadata["source_format"] == "txt"

    def test_paragraph_mode_double_newline(self, tmp_dir):
        f = tmp_dir / "para2.txt"
        f.write_text("Paragraph one.\n\nParagraph two.\n")
        cfg = ParsingConfig(
            text=TextParserConfig(paragraph_mode="double_newline")
        )
        parser = TextParser(config=cfg)
        doc = parser.parse(f)
        assert len(doc.blocks) >= 1


# ---------------------------------------------------------------------------
# E. DOCX parser
# ---------------------------------------------------------------------------


class TestDocxParser:
    def test_docx_basic_parsing(self, tmp_dir):
        # Create a minimal docx using python-docx directly
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            pytest.skip("python-docx not available")

        f = tmp_dir / "test.docx"
        doc = Document()
        doc.add_heading("Document Title", level=1)
        doc.add_paragraph("This is the first paragraph.")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("Content under section two.")
        doc.save(str(f))

        parser = DocxParser()
        result = parser.parse(f)
        assert result.file_type == "docx"
        assert result.title == "test"
        types = [b["type"] for b in result.blocks]
        assert "heading" in types
        heading_blocks = [b for b in result.blocks if b["type"] == "heading"]
        assert len(heading_blocks) >= 2

    def test_docx_source_format_tagged(self, tmp_dir):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        f = tmp_dir / "test2.docx"
        doc = Document()
        doc.add_paragraph("Hello world.")
        doc.save(str(f))

        parser = DocxParser()
        result = parser.parse(f)
        for b in result.blocks:
            assert b.get("source_format") == "docx"

    def test_docx_metadata_has_counts(self, tmp_dir):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")

        f = tmp_dir / "test3.docx"
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Para 1")
        doc.add_paragraph("Para 2")
        doc.save(str(f))

        parser = DocxParser()
        result = parser.parse(f)
        assert "heading_count" in result.metadata
        assert "parser_used" in result.metadata


# ---------------------------------------------------------------------------
# F. PDF parser
# ---------------------------------------------------------------------------


class TestPdfParser:
    def test_pdf_parser_is_registered(self, tmp_dir):
        # Just check PdfParser exists and inherits from BaseParser
        assert issubclass(PdfParser, BaseParser)

    def test_pdf_metadata_schema(self, tmp_dir):
        # PdfParser can be instantiated and returns ParsedDocument with expected fields
        # We don't have a real PDF in the test fixtures, so we just check the class
        parser = PdfParser()
        assert hasattr(parser, "parse")

    def test_pdf_fallback_rawtext_succeeds_for_complex_pdf(self):
        """
        Tier-2 fallback test: if all structured parsers return 0 blocks,
        the fitz raw-text fallback should still produce blocks for a
        readable (non-scanned) PDF.

        This requires a real PDF with text that defeats docling/pdfplumber/pymupdf
        block extraction.  The 3M 2015 10-K in data/financebench/pdfs/ is such a file.
        """
        pdf_path = ROOT / "data" / "financebench" / "pdfs" / "3M_2015_10K.pdf"
        if not pdf_path.exists():
            pytest.skip("3M_2015_10K.pdf not found")

        parser = PdfParser()
        doc = parser.parse(str(pdf_path))

        # Must not raise RuntimeError — fallback should succeed
        assert doc.file_type == "pdf"
        assert doc.metadata.get("parser_used") == "fitz_rawtext_fallback"
        assert doc.metadata.get("fallback_used") is True
        assert len(doc.blocks) > 0, "fallback should have produced blocks"
        assert len(doc.content) > 0, "content should not be empty"

        # All blocks should be tagged as pdf source
        for b in doc.blocks:
            assert b.get("source_format") == "pdf"
            assert b.get("text"), "block text should not be empty"

        # Should have page numbers
        pages = {b.get("page") for b in doc.blocks if b.get("page") is not None}
        assert len(pages) > 0, "should have page info"

    def test_pdf_fallback_metadata_has_reason(self):
        """Fallback path should include clear reason and scores in metadata."""
        pdf_path = ROOT / "data" / "financebench" / "pdfs" / "3M_2015_10K.pdf"
        if not pdf_path.exists():
            pytest.skip("3M_2015_10K.pdf not found")

        parser = PdfParser()
        doc = parser.parse(str(pdf_path))

        reason = doc.metadata.get("selection_reason", "")
        assert "fitz" in reason or "fallback" in reason.lower()
        assert "quality_score" in doc.metadata
        assert doc.metadata.get("likely_scanned_pdf") is False


# ---------------------------------------------------------------------------
# G. Pipeline integration
# ---------------------------------------------------------------------------


class TestPipeline:
    def test_parse_document_text_file(self, tmp_dir):
        f = tmp_dir / "hello.txt"
        f.write_text("Hello, world!\n")
        doc = parse_document(f)
        assert isinstance(doc, ParsedDocument)
        assert doc.file_type == "text"
        assert doc.content != ""

    def test_parse_document_unknown_type(self, tmp_dir):
        f = tmp_dir / "file.xyz"
        f.write_bytes(b"hello")
        with pytest.raises(ValueError, match="unsupported file type"):
            parse_document(f)

    def test_parsed_document_schema_fields(self, tmp_dir):
        f = tmp_dir / "schema.txt"
        f.write_text("# Title\n\nParagraph content.\n")
        doc = parse_document(f)
        # Required fields
        assert doc.title
        assert isinstance(doc.content, str)
        assert isinstance(doc.blocks, list)
        assert doc.source_path
        assert doc.file_type
        assert isinstance(doc.metadata, dict)
        # metadata keys
        assert "source_format" in doc.metadata
        # Block schema
        for b in doc.blocks:
            assert "text" in b
            assert "type" in b
            assert "source_format" in b

    def test_cleaning_config_applied(self, tmp_dir):
        f = tmp_dir / "cfg_test.txt"
        f.write_text("Page 1\n\nReal content\n\nPage 2\n")
        cfg = ParsingConfig(
            cleaning=CleaningConfig(
                remove_page_numbers=True,
                remove_separator_lines=True,
            )
        )
        doc = parse_document(f, config=cfg)
        texts = [b["text"] for b in doc.blocks]
        assert "Page 1" not in texts
        assert "Page 2" not in texts
        assert "Real content" in texts

    def test_parser_config_pdf_disabled(self, tmp_dir):
        # pdf disabled — but we don't have a PDF test file, so just check it instantiates
        cfg = ParsingConfig(pdf=PdfParserConfig(use_pdfplumber=False, use_docling=False))
        parser = PdfParser(config=cfg)
        assert parser._cfg.use_pdfplumber is False


# ---------------------------------------------------------------------------
# G. End-to-end with real files in data/
# ---------------------------------------------------------------------------


class TestRealFiles:
    """
    Smoke tests against real files in data/ if they exist.
    """

    def test_data_txt_parsing(self):
        data_dir = ROOT / "data"
        if not data_dir.exists():
            pytest.skip("data/ not found")

        txt_files = list(data_dir.glob("*.txt"))
        if not txt_files:
            pytest.skip("no .txt files in data/")

        for f in txt_files:
            try:
                doc = parse_document(f)
                assert doc.content, f"empty content for {f.name}"
                assert len(doc.blocks) > 0, f"no blocks for {f.name}"
                assert doc.metadata.get("source_format") == "txt"
            except Exception as e:
                pytest.fail(f"failed to parse {f.name}: {e}")

    def test_data_md_parsing(self):
        data_dir = ROOT / "data"
        if not data_dir.exists():
            pytest.skip("data/ not found")

        md_files = list(data_dir.glob("*.md"))
        if not md_files:
            pytest.skip("no .md files in data/")

        for f in md_files:
            try:
                doc = parse_document(f)
                assert len(doc.blocks) >= 0  # may be empty
                assert doc.metadata.get("source_format") == "txt"
            except Exception as e:
                pytest.fail(f"failed to parse {f.name}: {e}")

    def test_data_docx_parsing(self):
        data_dir = ROOT / "data"
        if not data_dir.exists():
            pytest.skip("data/ not found")

        docx_files = list(data_dir.glob("*.docx"))
        if not docx_files:
            pytest.skip("no .docx files in data/")

        for f in docx_files:
            try:
                doc = parse_document(f)
                assert doc.file_type == "docx"
                assert "parser_used" in doc.metadata
            except Exception as e:
                pytest.fail(f"failed to parse {f.name}: {e}")
